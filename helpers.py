import os
import re
import io
import json
import time
import tempfile
import pytesseract
import subprocess
from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document as DocxDocument
from fpdf import FPDF
from pdf2image import convert_from_bytes, convert_from_path

def clean_cv_text(text):
    if not text:
        return ""

    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def is_text_low_quality(text):
    if not text:
        return True

    text = clean_cv_text(text)
    words = text.split()

    if len(words) < 80:
        return True

    alpha_chars = sum(c.isalpha() for c in text)
    alpha_ratio = alpha_chars / max(len(text), 1)

    if alpha_ratio < 0.45:
        return True

    sections = ["experience", "education", "skills", "projects", "summary", "profile"]

    found_sections = sum(
        1 for section in sections
        if section in text.lower()
    )

    return found_sections < 2


def ocr_pdf_from_bytes(file_bytes, max_pages=3):
    text = ""

    images = convert_from_bytes(
        file_bytes,
        first_page=1,
        last_page=max_pages,
        dpi=200
    )

    for image in images:
        text += pytesseract.image_to_string(image) + "\n"

    return clean_cv_text(text)


def ocr_pdf_from_path(file_path, max_pages=3):
    text = ""

    images = convert_from_path(
        file_path,
        first_page=1,
        last_page=max_pages,
        dpi=200
    )

    for image in images:
        text += pytesseract.image_to_string(image) + "\n"

    return clean_cv_text(text)


def extract_text_from_file(file):
    try:
        if not file:
            return ""

        filename = (getattr(file, "filename", "") or "").lower()
        file.seek(0)

        if filename.endswith(".pdf"):
            reader = PdfReader(file)
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        if filename.endswith(".docx"):
            file_bytes = file.read()
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            try:
                return extract_docx(tmp_path)
            finally:
                safe_remove(tmp_path)

        if filename.endswith(".doc"):
            file_bytes = file.read()
            with tempfile.TemporaryDirectory() as tmpdir:
                doc_path = os.path.join(tmpdir, "input.doc")
                with open(doc_path, "wb") as f:
                    f.write(file_bytes)

                return (
                    try_libreoffice(doc_path, tmpdir)
                    or try_antiword(doc_path)
                    or try_binary_doc(file_bytes)
                )

        return ""

    except Exception as e:
        print("❌ Extraction error:", e)
        return ""


# def extract_text_from_path(file_path):
#     try:
#         path = os.fspath(file_path)
#         ext = os.path.splitext(path.lower())[1]

#         if ext == ".pdf":
#             reader = PdfReader(path)
#             return "\n".join(page.extract_text() or "" for page in reader.pages)

#         if ext == ".docx":
#             return extract_docx(path)

#         if ext == ".doc":
#             with open(path, "rb") as f:
#                 file_bytes = f.read()

#             tmpdir = os.path.dirname(path) or tempfile.gettempdir()
#             return (
#                 try_libreoffice(path, tmpdir)
#                 or try_antiword(path)
#                 or try_binary_doc(file_bytes)
#             )

#         print("❌ Unsupported file type:", path)
#         return ""

#     except Exception as e:
#         print("❌ Extraction error:", e)
#         return ""


def extract_text_from_path(file_path):
    try:
        path = os.fspath(file_path)
        ext = os.path.splitext(path.lower())[1]
        text = ""

        if ext == ".pdf":
            reader = PdfReader(path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            text = clean_cv_text(text)

            if is_text_low_quality(text):
                print(f"⚠️ OCR fallback used for ATS PDF: {path}")
                ocr_text = ocr_pdf_from_path(path)

                if len(ocr_text.split()) > len(text.split()):
                    text = ocr_text

            return clean_cv_text(text)

        if ext == ".docx":
            text = extract_docx(path)
            return clean_cv_text(text)

        if ext == ".doc":
            with open(path, "rb") as f:
                file_bytes = f.read()

            tmpdir = os.path.dirname(path) or tempfile.gettempdir()

            text = (
                try_libreoffice(path, tmpdir)
                or try_antiword(path)
                or try_binary_doc(file_bytes)
            )

            return clean_cv_text(text)

        print("❌ Unsupported file type:", path)
        return ""

    except Exception as e:
        print("❌ Extraction error:", e)

        try:
            path = os.fspath(file_path)
            ext = os.path.splitext(path.lower())[1]

            if ext == ".pdf":
                print(f"⚠️ Extraction failed, trying OCR: {path}")
                return clean_cv_text(ocr_pdf_from_path(path))

        except Exception as ocr_error:
            print("OCR fallback failed:", str(ocr_error))

        return ""
    


def extract_docx(source):
    doc = DocxDocument(source)

    para_text = "\n".join([p.text for p in doc.paragraphs]).strip()

    table_lines = []
    seen = set()

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                line = " | ".join(cells)
                if line not in seen:
                    seen.add(line)
                    table_lines.append(line)

    table_text = "\n".join(table_lines)

    if len(para_text) > 100:
        extra = [line for line in table_lines if line[:50].lower() not in para_text.lower()]
        return para_text + ("\n\n" + "\n".join(extra) if extra else "")

    return table_text or para_text


# def try_libreoffice(doc_path, tmpdir):
#     try:
#         subprocess.run(
#             ["libreoffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", tmpdir],
#             capture_output=True,
#             timeout=30,
#         )
#         docx_path = os.path.join(tmpdir, "input.docx")
#         if os.path.exists(docx_path):
#             text = extract_docx(docx_path)
#             if text and len(text.strip()) > 50:
#                 return text
#     except Exception:
#         pass

#     return ""

def try_libreoffice(doc_path, tmpdir):
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", tmpdir],
            capture_output=True,
            timeout=30,
        )

        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        docx_path = os.path.join(tmpdir, base_name + ".docx")

        if os.path.exists(docx_path):
            text = extract_docx(docx_path)

            if text and len(text.strip()) > 50:
                return text

    except Exception:
        pass

    return ""


# def extract_text_from_bytes(filename, file_bytes):
#     text = ""
#     try:
#         # PDF
#         if filename.lower().endswith(".pdf"):
#             reader = PdfReader(io.BytesIO(file_bytes))
#             for page in reader.pages:
#                 text += (page.extract_text() or "") + "\n"

#         # DOCX
#         elif filename.lower().endswith(".docx"):

#             doc = DocxDocument(io.BytesIO(file_bytes))

#             for para in doc.paragraphs:
#                 text += para.text + "\n"

#         # TXT fallback
#         else:
#             text = file_bytes.decode("utf-8", errors="ignore")

#     except Exception as e:
#         print("extract_text_from_bytes Error:", str(e))

#     return text.strip()


def extract_text_from_bytes(filename, file_bytes):
    text = ""
    filename_lower = filename.lower()

    try:
        if filename_lower.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(file_bytes))

            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"

            text = clean_cv_text(text)

            if is_text_low_quality(text):
                print(f"⚠️ OCR fallback used for manual PDF: {filename}")
                ocr_text = ocr_pdf_from_bytes(file_bytes)

                if len(ocr_text.split()) > len(text.split()):
                    text = ocr_text

        elif filename_lower.endswith(".docx"):
            doc = DocxDocument(io.BytesIO(file_bytes))

            for para in doc.paragraphs:
                text += para.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"

            text = clean_cv_text(text)

        else:
            text = file_bytes.decode("utf-8", errors="ignore")
            text = clean_cv_text(text)

    except Exception as e:
        print("extract_text_from_bytes Error:", str(e))

        if filename_lower.endswith(".pdf"):
            try:
                print(f"⚠️ Normal extraction failed, trying OCR: {filename}")
                text = ocr_pdf_from_bytes(file_bytes)
            except Exception as ocr_error:
                print("OCR fallback failed:", str(ocr_error))

    return clean_cv_text(text)


def try_antiword(doc_path):
    try:
        result = subprocess.run(
            ["antiword", doc_path],
            capture_output=True,
            text=True,
            timeout=15,
        )

        if result.returncode == 0 and len(result.stdout.strip()) > 50:
            return result.stdout.strip()

    except Exception:
        pass

    return ""


def try_binary_doc(file_bytes):
    runs = []
    current = []

    for i in range(0, len(file_bytes) - 1, 2):
        lo, hi = file_bytes[i], file_bytes[i + 1]

        if hi == 0 and (32 <= lo <= 126 or lo in (9, 10, 13)):
            current.append(chr(lo))
        else:
            if len(current) > 20:
                runs.append("".join(current).strip())
            current = []

    if len(current) > 20:
        runs.append("".join(current).strip())

    runs.sort(key=len, reverse=True)

    clean = [
        r for r in runs[:15]
        if sum(1 for c in r if c.isalpha()) / max(len(r), 1) > 0.4 and len(r) > 30
    ]

    result = "\n".join(clean)
    return result if len(result) > 50 else ""


# def clean_json(raw):
#     raw = re.sub(r"```json\s*", "", raw)
#     raw = re.sub(r"```\s*", "", raw)

#     match = re.search(r"[\{\[].*[\}\]]", raw, re.DOTALL)
#     if match:
#         raw = match.group(0)

#     raw = raw.replace("\n", " ").replace("\r", "")
#     raw = re.sub(r",\s*}", "}", raw)
#     raw = re.sub(r",\s*]", "]", raw)
#     raw = re.sub(r"[\x00-\x1f]", " ", raw)

#     try:
#         data = json.loads(raw)
#     except json.JSONDecodeError:
#         depth = 0
#         start = raw.index("{")

#         for i in range(start, len(raw)):
#             if raw[i] == "{":
#                 depth += 1
#             elif raw[i] == "}":
#                 depth -= 1

#             if depth == 0:
#                 data = json.loads(raw[start:i + 1])
#                 break
#         else:
#             raise ValueError("Could not parse JSON from AI response")

#     return data[0] if isinstance(data, list) else data


def clean_json(text):
    if not text:
        raise ValueError("Empty AI response")

    text = text.strip()

    # Strip markdown code fences
    text = re.sub(r"```json\s*", "", text)
    text = re.sub(r"```\s*", "", text)
    text = text.strip()

    # Replace smart/curly quotes with straight quotes
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u00ab", '"').replace("\u00bb", '"')

    # Extract the JSON object
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        raise ValueError("No JSON object found in AI response")

    cleaned = match.group(0)

    # Remove trailing commas before } or ]
    cleaned = re.sub(r",\s*}", "}", cleaned)
    cleaned = re.sub(r",\s*]", "]", cleaned)

    # Remove control characters (except newline/tab) that break JSON
    cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", cleaned)

    # First attempt: try parsing as-is
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Second attempt: collapse newlines inside string values
    # Replace actual newlines with escaped \\n (within strings)
    collapsed = cleaned.replace("\r\n", "\\n").replace("\r", "\\n").replace("\n", "\\n")
    # Fix double-escaped newlines
    collapsed = collapsed.replace("\\\\n", "\\n")
    try:
        return json.loads(collapsed)
    except json.JSONDecodeError:
        pass

    # Third attempt: try to fix single quotes used as string delimiters
    # Only do this if there are no double-quoted strings
    if cleaned.count("'") > cleaned.count('"'):
        single_fixed = cleaned.replace("'", '"')
        single_fixed = re.sub(r",\s*}", "}", single_fixed)
        single_fixed = re.sub(r",\s*]", "]", single_fixed)
        try:
            return json.loads(single_fixed)
        except json.JSONDecodeError:
            pass

    # Fourth attempt: aggressive cleanup — strip everything to one line
    oneline = re.sub(r"\s+", " ", cleaned)
    oneline = re.sub(r",\s*}", "}", oneline)
    oneline = re.sub(r",\s*]", "]", oneline)
    try:
        return json.loads(oneline)
    except json.JSONDecodeError:
        pass

    # Final fallback: delegate to the flexible parser
    try:
        return clean_json_flexible(text)
    except Exception:
        pass

    raise ValueError(f"Could not parse AI response as JSON")

def clean_json_flexible(raw):
    raw = re.sub(r"```json\s*", "", raw)
    raw = re.sub(r"```\s*", "", raw)
    raw = raw.strip()

    try:
        return json.loads(raw)
    except Exception:
        pass

    cleaned = raw.replace("\n", " ").replace("\r", "")
    cleaned = re.sub(r",\s*}", "}", cleaned)
    cleaned = re.sub(r",\s*]", "]", cleaned)
    cleaned = re.sub(r"[\x00-\x1f]", " ", cleaned)

    try:
        return json.loads(cleaned)
    except Exception:
        pass

    arr_match = re.search(r"\[.*\]", cleaned, re.DOTALL)
    if arr_match:
        try:
            return json.loads(arr_match.group(0))
        except Exception:
            pass

    obj_match = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if obj_match:
        try:
            return json.loads(obj_match.group(0))
        except Exception:
            pass

    questions = []
    pattern = r'\{"question"\s*:\s*"([^"]+)"\s*,\s*"options"\s*:\s*\[([^\]]+)\]\s*\}'

    for match in re.finditer(pattern, cleaned):
        opts = [o.strip().strip('"').strip("'") for o in match.group(2).split(",")]
        questions.append({
            "question": match.group(1),
            "options": opts[:4],
        })

    if questions:
        return questions

    raise ValueError("Could not parse AI response as JSON")


def clean_html_to_text(html):
    soup = BeautifulSoup(html or "", "html.parser")
    return soup.get_text(separator="\n")


def safe_remove(path):
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except OSError:
        pass


def create_jd_docx(html_content):
    doc = DocxDocument()

    text = html_content or ""
    text = re.sub(r"<h3[^>]*>", "\n[HEADING]", text)
    text = re.sub(r"</h3>", "[/HEADING]\n", text)
    text = re.sub(r"<li[^>]*>", "\n• ", text)
    text = re.sub(r"<strong[^>]*>", "[B]", text)
    text = re.sub(r"</strong>", "[/B]", text)
    text = re.sub(r"<mark[^>]*>", "", text)
    text = re.sub(r"</mark>", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        if "[HEADING]" in line:
            heading_text = (
                line.replace("[HEADING]", "")
                .replace("[/HEADING]", "")
                .replace("[B]", "")
                .replace("[/B]", "")
                .strip()
            )
            doc.add_heading(heading_text, level=2)

        elif line.startswith("•"):
            item_text = line[1:].strip().replace("[B]", "").replace("[/B]", "")
            doc.add_paragraph(item_text, style="List Bullet")

        else:
            paragraph = doc.add_paragraph()
            parts = re.split(r"(\[B\].*?\[/B\])", line)

            for part in parts:
                if part.startswith("[B]"):
                    run = paragraph.add_run(part.replace("[B]", "").replace("[/B]", ""))
                    run.bold = True
                elif part.strip():
                    paragraph.add_run(part)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_jd_pdf(html_content):
    text = html_content or ""

    text = re.sub(r"<h3[^>]*>", "\n[HEADING]", text)
    text = re.sub(r"</h3>", "[/HEADING]\n", text)
    text = re.sub(r"<li[^>]*>", "\n[BULLET]", text)
    text = re.sub(r"<strong[^>]*>", "[B]", text)
    text = re.sub(r"</strong>", "[/B]", text)
    text = re.sub(r"<mark[^>]*>", "", text)
    text = re.sub(r"</mark>", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    replacements = {
        "\u2022": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
        "\u00a0": " ",
        "\u2023": "-",
        "\u25cf": "-",
        "\u25cb": "-",
        "\u2010": "-",
        "\u2011": "-",
        "\u2012": "-",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", "", 11)

    for line in text.split("\n"):
        line = line.strip()

        if not line:
            pdf.ln(4)
            continue

        if "[HEADING]" in line:
            heading = (
                line.replace("[HEADING]", "")
                .replace("[/HEADING]", "")
                .replace("[B]", "")
                .replace("[/B]", "")
                .strip()
            )

            pdf.ln(6)
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 8, heading, new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(37, 99, 235)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(4)
            pdf.set_font("Helvetica", "", 11)

        elif "[BULLET]" in line:
            item = line.replace("[BULLET]", "").replace("[B]", "").replace("[/B]", "").strip()
            pdf.set_x(15)
            pdf.cell(5, 6, "-", new_x="END")
            pdf.multi_cell(170, 6, item, new_x="LMARGIN", new_y="NEXT")

        else:
            parts = re.split(r"(\[B\].*?\[/B\])", line)

            if any("[B]" in part for part in parts):
                for part in parts:
                    if part.startswith("[B]"):
                        pdf.set_font("Helvetica", "B", 11)
                        pdf.write(6, part.replace("[B]", "").replace("[/B]", ""))
                        pdf.set_font("Helvetica", "", 11)
                    elif part.strip():
                        pdf.write(6, part)
                pdf.ln(6)
            else:
                pdf.multi_cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")

    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer


def create_cv_report(results):
    doc = DocxDocument()

    doc.add_heading("Candidate Screening Report", level=1)
    doc.add_paragraph(f"Generated on {time.strftime('%d %B %Y')} • {len(results)} Candidates Screened")
    doc.add_paragraph("─" * 60)

    for data in results:
        score = data.get("overallScore", 0)
        name = data.get("candidate_name", "Unknown")
        rec = data.get("recommendation", "")
        rationale = data.get("rationale", "")

        doc.add_heading(f"{name}  —  {score}%", level=2)
        doc.add_paragraph(f"Phone: {data.get('phone_number', 'Not found')}")
        doc.add_paragraph(f"Recommendation: {rec}")
        doc.add_paragraph(f"Rationale: {rationale}")

        doc.add_heading("Strengths", level=3)
        for key in ["NIRF_and_Pedigree", "Experience_Alignment", "Projects_and_Quantifiable_Impact"]:
            items = data.get("strengths", {}).get(key, [])
            for item in items:
                if item and item.strip().lower() != "none":
                    doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("Proximity", level=3)
        for item in data.get("proximity_matches", []):
            if item and item.strip().lower() != "none":
                doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("Gaps", level=3)
        for key in ["Functional_Gaps", "Domain_Mismatch"]:
            items = data.get("gaps", {}).get(key, [])
            for item in items:
                if item and item.strip().lower() != "none":
                    doc.add_paragraph(item, style="List Bullet")

        doc.add_paragraph("─" * 60)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer